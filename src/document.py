"""
Programmatic Word document builder for resumes and cover letters.

Formatting spec:
    - Font: Arial throughout (default / CLASSIC theme)
    - Body text: 10pt
    - Section headings: 12pt, Bold, ALL CAPS
    - Name header: 14pt, Bold, Centered
    - Contact line: 10pt, Centered
    - Company/location line: 10pt, Bold — company LEFT, dates RIGHT-TABBED
    - Role: 10pt, Italic
    - Bullets: 10pt, 0.25" hanging indent
    - Education & Certs: borderless 2-column table (left=education, right=certs);
        single-column if no certs returned
"""

import os
import re
from datetime import date
from typing import TYPE_CHECKING

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips

from src.models import CoverLetterJSON, ResumeJSON
from src.themes import ThemeConfig, CLASSIC

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

RIGHT_TAB_POS = Twips(9360)  # 6.5" at 1" margins on US Letter


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------


def _set_font(run, size_pt: int, bold: bool = False, italic: bool = False,
                color: RGBColor | None = None, font_name: str = "Arial"):
    # pylint: disable=too-many-arguments,too-many-positional-arguments
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def _add_right_tab(paragraph, pos_twips: int | None = None):
    """Add a right-aligned tab stop at the right margin."""
    paragraph_properties = paragraph._p.get_or_add_pPr()  # pylint: disable=protected-access
    tabs_element = OxmlElement("w:tabs")
    tab_element = OxmlElement("w:tab")
    tab_element.set(qn("w:val"), "right")
    tab_element.set(qn("w:pos"), str(pos_twips or RIGHT_TAB_POS))
    tabs_element.append(tab_element)
    paragraph_properties.append(tabs_element)


def _set_space(paragraph, before_pt: float = 0, after_pt: float = 0,
                line_rule: str | None = None, line_val: int | None = None):
    paragraph_properties = paragraph._p.get_or_add_pPr()  # pylint: disable=protected-access
    spacing_element = OxmlElement("w:spacing")
    spacing_element.set(qn("w:before"), str(int(before_pt * 20)))
    spacing_element.set(qn("w:after"), str(int(after_pt * 20)))
    if line_rule and line_val:
        spacing_element.set(qn("w:line"), str(line_val))
        spacing_element.set(qn("w:lineRule"), line_rule)
    paragraph_properties.append(spacing_element)


def _no_border_cell(cell):
    """Remove all borders from a table cell."""
    tc_obj = cell._tc  # pylint: disable=protected-access
    tc_pr = tc_obj.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tc_borders.append(el)
    tc_pr.append(tc_borders)


def _shade_cell(cell, hex_color: str):
    """Set cell background fill color via XML."""
    tc_obj = cell._tc  # pylint: disable=protected-access
    cell_properties = tc_obj.get_or_add_tcPr()
    shading_element = OxmlElement("w:shd")
    shading_element.set(qn("w:val"), "clear")
    shading_element.set(qn("w:fill"), hex_color)
    cell_properties.append(shading_element)


def _cell_bullet(cell, text: str):
    """Add a bullet-style paragraph inside a table cell."""
    paragraph = cell.add_paragraph()
    _set_space(paragraph, before_pt=1, after_pt=1)
    # Hanging indent to mimic bullet style
    paragraph_properties = paragraph._p.get_or_add_pPr()  # pylint: disable=protected-access
    indent_element = OxmlElement("w:ind")
    indent_element.set(qn("w:left"), "360")   # 0.25"
    indent_element.set(qn("w:hanging"), "360")
    paragraph_properties.append(indent_element)
    # Bullet character + text
    bullet_run = paragraph.add_run("\u2022\t")
    _set_font(bullet_run, 10)
    text_run = paragraph.add_run(text)
    _set_font(text_run, 10)
    return paragraph


def _section_heading(doc: "DocxDocument", text: str, theme: "ThemeConfig"):
    paragraph = doc.add_paragraph()
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.LEFT if theme.name_align == "left"
        else WD_ALIGN_PARAGRAPH.CENTER
    )
    _set_space(paragraph, before_pt=8, after_pt=2)
    heading_run = paragraph.add_run(text.upper())
    _set_font(heading_run, int(theme.heading_pt), bold=True, color=RGBColor(*theme.accent_color),
                font_name=theme.font)
    paragraph_properties = paragraph._p.get_or_add_pPr()  # pylint: disable=protected-access
    if theme.heading_rule:
        border_element = OxmlElement("w:pBdr")
        top_border = OxmlElement("w:top")
        top_border.set(qn("w:val"), "single")
        top_border.set(qn("w:sz"), "4")
        top_border.set(qn("w:space"), "1")
        top_border.set(qn("w:color"), "000000")
        border_element.append(top_border)
        paragraph_properties.append(border_element)
    if theme.heading_underline:
        border_element = OxmlElement("w:pBdr")
        bottom_border = OxmlElement("w:bottom")
        bottom_border.set(qn("w:val"), "single")
        bottom_border.set(qn("w:sz"), "6")
        bottom_border.set(qn("w:space"), "1")
        red, green, blue = theme.accent_color
        bottom_border.set(qn("w:color"), f"{red:02x}{green:02x}{blue:02x}")
        border_element.append(bottom_border)
        paragraph_properties.append(border_element)
    keep_next = OxmlElement("w:keepNext")
    paragraph_properties.append(keep_next)
    return paragraph


def _body_paragraph(doc: "DocxDocument", text: str, italic: bool = False,
                    before_pt: float = 0, after_pt: float = 2, body_pt: int = 10,
                    font_name: str = "Arial"):
    # pylint: disable=too-many-arguments,too-many-positional-arguments
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_space(paragraph, before_pt=before_pt, after_pt=after_pt)
    text_run = paragraph.add_run(text)
    _set_font(text_run, body_pt, italic=italic, font_name=font_name)
    return paragraph


def _bullet(doc: "DocxDocument", text: str, body_pt: int = 10, font_name: str = "Arial"):
    paragraph = doc.add_paragraph(style="List Bullet")
    _set_space(paragraph, before_pt=0, after_pt=1)
    # Clear default run and add our own with correct font
    for run in paragraph.runs:
        paragraph._p.remove(run._r)  # pylint: disable=protected-access
    text_run = paragraph.add_run(text)
    _set_font(text_run, body_pt, font_name=font_name)
    return paragraph


def _company_line(doc: "DocxDocument", company: str, location: str, dates: str,
                    right_tab_twips: int | None = None, body_pt: int = 10,
                    font_name: str = "Arial"):
    """Company + location LEFT, dates RIGHT using a tab stop."""
    # pylint: disable=too-many-arguments,too-many-positional-arguments
    paragraph = doc.add_paragraph()
    _set_space(paragraph, before_pt=6, after_pt=0)
    if right_tab_twips:
        _add_right_tab(paragraph, right_tab_twips)
    else:
        _add_right_tab(paragraph)
    left_text = f"{company}  |  {location}" if location else company
    run_left = paragraph.add_run(left_text)
    _set_font(run_left, body_pt, bold=True, font_name=font_name)
    run_tab = paragraph.add_run("\t")
    run_tab.font.size = Pt(body_pt)
    run_dates = paragraph.add_run(dates)
    _set_font(run_dates, body_pt, bold=True, font_name=font_name)
    return paragraph


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _end_year(experience_entry) -> int:
    """Extract the end year from an ExperienceEntry's dates string for sorting.

    Returns 9999 for 'Present' / 'Current' so active roles sort to the top.
    Falls back to the last 4-digit year found, or 0 if none.
    """
    dates = experience_entry.dates or ""
    if re.search(r"\b(present|current)\b", dates, re.IGNORECASE):
        return 9999
    years = re.findall(r"\b(19|20)\d{2}\b", dates)
    return int(years[-1]) if years else 0


def _render_sidebar_column(left_cell, resume_json, personal, education, theme: ThemeConfig):
    """Render the left sidebar column of the Creative resume."""
    # pylint: disable=too-many-locals,too-many-statements
    # --- Sidebar Header (Name) ---
    name_paragraph = left_cell.paragraphs[0]
    _set_space(name_paragraph, before_pt=6, after_pt=2)
    name_run = name_paragraph.add_run(personal["name"])
    _set_font(name_run, int(theme.name_pt), bold=True, color=RGBColor(255, 255, 255),
                font_name=theme.font)

    # --- Location & Contact ---
    loc = personal.get("location", {})
    location_str = (
        ", ".join(part for part in [loc.get("city", ""), loc.get("region", "")] if part)
        if isinstance(loc, dict) else str(loc or "")
    )
    for line in filter(None, [
        location_str,
        personal.get("phone", ""),
        personal.get("email", ""),
    ]):
        contact_paragraph = left_cell.add_paragraph()
        _set_space(contact_paragraph, before_pt=1, after_pt=1)
        contact_run = contact_paragraph.add_run(line)
        _set_font(contact_run, int(theme.body_pt) - 1, color=RGBColor(200, 200, 200),
                    font_name=theme.font)

    def __sidebar_section(text: str):
        section_paragraph = left_cell.add_paragraph()
        _set_space(section_paragraph, before_pt=10, after_pt=3)
        section_run = section_paragraph.add_run(text.upper())
        _set_font(section_run, int(theme.heading_pt), bold=True, color=RGBColor(180, 180, 180),
                    font_name=theme.font)
        paragraph_properties = section_paragraph._p.get_or_add_pPr()  # pylint: disable=protected-access
        border_element = OxmlElement("w:pBdr")
        bottom_border = OxmlElement("w:bottom")
        bottom_border.set(qn("w:val"), "single")
        bottom_border.set(qn("w:sz"), "4")
        bottom_border.set(qn("w:space"), "1")
        bottom_border.set(qn("w:color"), "666666")
        border_element.append(bottom_border)
        paragraph_properties.append(border_element)

    # --- Skills ---
    __sidebar_section("Skills")
    for category in resume_json.skill_categories:
        label_paragraph = left_cell.add_paragraph()
        _set_space(label_paragraph, before_pt=3, after_pt=1)
        label_run = label_paragraph.add_run(category.name)
        _set_font(label_run, int(theme.body_pt) - 1, bold=True, color=RGBColor(220, 220, 220),
                    font_name=theme.font)
        for skill in category.skills:
            skill_paragraph = left_cell.add_paragraph()
            _set_space(skill_paragraph, before_pt=1, after_pt=0)
            skill_run = skill_paragraph.add_run(skill)
            _set_font(skill_run, int(theme.body_pt) - 1, color=RGBColor(190, 190, 190),
                        font_name=theme.font)

    # --- Education ---
    if education:
        __sidebar_section("Education")
        for ed in education:
            degree_parts = [ed.get("studyType", ""), ed.get("area", "")]
            degree = " - ".join(part for part in degree_parts if part)
            edu_paragraph = left_cell.add_paragraph()
            _set_space(edu_paragraph, before_pt=3, after_pt=1)
            edu_run = edu_paragraph.add_run(degree or ed.get("institution", ""))
            _set_font(edu_run, int(theme.body_pt) - 1, bold=True, color=RGBColor(220, 220, 220),
                        font_name=theme.font)
            if degree and ed.get("institution"):
                inst_paragraph = left_cell.add_paragraph()
                _set_space(inst_paragraph, before_pt=0, after_pt=0)
                inst_run = inst_paragraph.add_run(ed["institution"])
                _set_font(inst_run, int(theme.body_pt) - 1, color=RGBColor(190, 190, 190),
                            font_name=theme.font)

    # --- Certifications ---
    if resume_json.certifications:
        __sidebar_section("Certifications")
        for cert in resume_json.certifications:
            cert_paragraph = left_cell.add_paragraph()
            _set_space(cert_paragraph, before_pt=2, after_pt=1)
            cert_run = cert_paragraph.add_run(cert)
            _set_font(cert_run, int(theme.body_pt) - 1, color=RGBColor(190, 190, 190),
                        font_name=theme.font)


def _render_main_column(right_cell, resume_json, main_width_emu: int, theme: ThemeConfig):
    """Render the right main column of the Creative resume."""
    # pylint: disable=too-many-locals,too-many-statements
    def __main_section(text: str):
        section_paragraph = right_cell.add_paragraph()
        _set_space(section_paragraph, before_pt=10, after_pt=3)
        section_run = section_paragraph.add_run(text.upper())
        _set_font(section_run, int(theme.heading_pt), bold=True, font_name=theme.font)
        paragraph_properties = section_paragraph._p.get_or_add_pPr()  # pylint: disable=protected-access
        if theme.heading_underline:
            border_element = OxmlElement("w:pBdr")
            bottom_border = OxmlElement("w:bottom")
            bottom_border.set(qn("w:val"), "single")
            bottom_border.set(qn("w:sz"), "6")
            bottom_border.set(qn("w:space"), "1")
            bottom_border.set(qn("w:color"), "000000")
            border_element.append(bottom_border)
            paragraph_properties.append(border_element)
        keep_next = OxmlElement("w:keepNext")
        paragraph_properties.append(keep_next)

    right_tab_stop = int(main_width_emu / 635)

    __main_section("Professional Summary")
    summary_paragraph = right_cell.add_paragraph()
    _set_space(summary_paragraph, before_pt=2, after_pt=4)
    summary_run = summary_paragraph.add_run(resume_json.summary.strip())
    _set_font(summary_run, int(theme.body_pt), font_name=theme.font)

    __main_section("Professional Experience")
    experience_list = sorted(resume_json.experience, key=_end_year, reverse=True)
    for exp in experience_list:
        company_paragraph = right_cell.add_paragraph()
        _set_space(company_paragraph, before_pt=6, after_pt=0)
        if right_tab_stop:
            _add_right_tab(company_paragraph, right_tab_stop)
        company_run = company_paragraph.add_run(f"{exp.company}")
        _set_font(company_run, int(theme.body_pt), bold=True, font_name=theme.font)
        dates_run = company_paragraph.add_run("\t" + exp.dates)
        _set_font(dates_run, int(theme.body_pt), bold=True, font_name=theme.font)
        role_paragraph = right_cell.add_paragraph()
        _set_space(role_paragraph, before_pt=1, after_pt=1)
        role_run = role_paragraph.add_run(exp.role)
        _set_font(role_run, int(theme.body_pt), italic=True, font_name=theme.font)
        for bullet in exp.bullets:
            bullet_paragraph = right_cell.add_paragraph(style="List Bullet")
            _set_space(bullet_paragraph, before_pt=0, after_pt=1)
            for run in bullet_paragraph.runs:
                bullet_paragraph._p.remove(run._r)  # pylint: disable=protected-access
            bullet_run = bullet_paragraph.add_run(bullet)
            _set_font(bullet_run, int(theme.body_pt), font_name=theme.font)

    if resume_json.projects:
        __main_section(resume_json.projects_section_heading)
        for project in resume_json.projects:
            project_paragraph = right_cell.add_paragraph()
            _set_space(project_paragraph, before_pt=4, after_pt=0)
            project_title_run = project_paragraph.add_run(project.title)
            _set_font(project_title_run, int(theme.body_pt), bold=True, font_name=theme.font)
            for bullet in project.bullets:
                bullet_paragraph = right_cell.add_paragraph(style="List Bullet")
                _set_space(bullet_paragraph, before_pt=0, after_pt=1)
                for run in bullet_paragraph.runs:
                    bullet_paragraph._p.remove(run._r)  # pylint: disable=protected-access
                bullet_run = bullet_paragraph.add_run(bullet)
                _set_font(bullet_run, int(theme.body_pt), font_name=theme.font)


# ---------------------------------------------------------------------------
# Sidebar resume builder (Creative theme)
# ---------------------------------------------------------------------------


def _build_resume_sidebar(resume_json: ResumeJSON, personal: dict, education: list,
                            output_path: str, theme: ThemeConfig):
    """Render resume with two-column sidebar layout (Creative theme)."""
    # pylint: disable=too-many-locals
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(theme.margin_top)
        section.bottom_margin = Inches(theme.margin_bottom)
        section.left_margin = Inches(theme.margin_left)
        section.right_margin = Inches(theme.margin_right)

    section = doc.sections[0]
    total_width_emu = (
        (section.page_width or 0) -
        (section.left_margin or 0) -
        (section.right_margin or 0)
    )
    sidebar_width_emu = int(total_width_emu * 0.28)
    main_width_emu = total_width_emu - sidebar_width_emu

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 1)
    _no_border_cell(left_cell)
    _no_border_cell(right_cell)

    left_cell.width = sidebar_width_emu
    right_cell.width = main_width_emu
    red, green, blue = theme.sidebar_color
    _shade_cell(left_cell, f"{red:02x}{green:02x}{blue:02x}")

    _render_sidebar_column(left_cell, resume_json, personal, education, theme)

    _render_main_column(right_cell, resume_json, main_width_emu, theme)


    _dirname = os.path.dirname(output_path)
    if _dirname:
        os.makedirs(_dirname, exist_ok=True)
    doc.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# Resume builder
# ---------------------------------------------------------------------------


def _render_resume_header(doc: "DocxDocument", personal: dict, theme: ThemeConfig):
    # pylint: disable=too-many-locals
    name_paragraph = doc.add_paragraph()
    name_paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.LEFT if theme.name_align == "left"
        else WD_ALIGN_PARAGRAPH.CENTER
    )
    _set_space(name_paragraph, before_pt=0, after_pt=2)
    name_run = name_paragraph.add_run(personal["name"])
    _set_font(name_run, int(theme.name_pt), bold=True, color=RGBColor(*theme.accent_color),
                font_name=theme.font)

    loc = personal.get("location", {})
    loc_parts = [loc.get("city", ""), loc.get("region", ""), loc.get("countryCode", "")]
    location_str = (
        ", ".join(part for part in loc_parts if part) if isinstance(loc, dict) else str(loc or "")
    )
    linkedin = next(
        (profile.get("username", "") for profile in personal.get("profiles", [])
            if profile.get("network") == "LinkedIn"), ""
    )
    github = next(
        (profile.get("username", "") for profile in personal.get("profiles", [])
            if profile.get("network") == "GitHub"), ""
    )
    contact_parts = [
        location_str,
        personal.get("phone", ""),
        personal.get("email", ""),
        f"LinkedIn: {linkedin}" if linkedin else "",
        f"GitHub: {github}" if github else "",
    ]
    contact_line = " \u2219 ".join(part for part in contact_parts if part)
    contact_paragraph = doc.add_paragraph()
    contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_space(contact_paragraph, before_pt=0, after_pt=4)
    contact_run = contact_paragraph.add_run(contact_line)
    _set_font(contact_run, int(theme.body_pt), font_name=theme.font)


def _render_resume_experience(doc: "DocxDocument", experience, right_tab, theme: ThemeConfig):
    _section_heading(doc, "Professional Experience", theme)
    experience_list = sorted(experience, key=_end_year, reverse=True)
    for entry in experience_list:
        _company_line(doc, entry.company, "", entry.dates, right_tab_twips=right_tab,
                        body_pt=int(theme.body_pt), font_name=theme.font)
        _body_paragraph(doc, entry.role, italic=True, before_pt=1, after_pt=1,
                        body_pt=int(theme.body_pt), font_name=theme.font)
        for bullet in entry.bullets:
            _bullet(doc, bullet, body_pt=int(theme.body_pt), font_name=theme.font)


def _render_resume_skills(doc: "DocxDocument", skill_categories, theme: ThemeConfig):
    _section_heading(doc, "Skills", theme)
    for category in skill_categories:
        bullet_paragraph = doc.add_paragraph(style="List Bullet")
        _set_space(bullet_paragraph, before_pt=1, after_pt=1)
        for run in bullet_paragraph.runs:
            bullet_paragraph._p.remove(run._r)  # pylint: disable=protected-access
        label_run = bullet_paragraph.add_run(f"{category.name}: ")
        _set_font(label_run, int(theme.body_pt), bold=True, font_name=theme.font)
        skills_run = bullet_paragraph.add_run(", ".join(category.skills))
        _set_font(skills_run, int(theme.body_pt), font_name=theme.font)


def _render_resume_projects(doc: "DocxDocument", projects, heading, theme: ThemeConfig):
    _section_heading(doc, heading, theme)
    for project in projects:
        project_paragraph = doc.add_paragraph()
        _set_space(project_paragraph, before_pt=4, after_pt=0)
        title_run = project_paragraph.add_run(project.title)
        _set_font(title_run, int(theme.body_pt), bold=True, font_name=theme.font)
        if project.focus:
            focus_run = project_paragraph.add_run(f"  —  {project.focus}")
            _set_font(focus_run, int(theme.body_pt), italic=True, font_name=theme.font)
        for bullet in project.bullets:
            _bullet(doc, bullet, body_pt=int(theme.body_pt), font_name=theme.font)
        if project.url:
            url_paragraph = doc.add_paragraph()
            _set_space(url_paragraph, before_pt=0, after_pt=1)
            url_run = url_paragraph.add_run(project.url)
            _set_font(url_run, int(theme.body_pt), italic=True, font_name=theme.font)


def _render_resume_education_certs(doc: "DocxDocument", education, certifications,
                                    theme: ThemeConfig):
    has_certs = bool(certifications)
    heading_text = "Education & Certificates" if has_certs else "Education"
    _section_heading(doc, heading_text, theme)

    if has_certs:
        table = doc.add_table(rows=1, cols=2)
        table.autofit = True
        left_cell = table.cell(0, 0)
        right_cell = table.cell(0, 1)
        _no_border_cell(left_cell)
        _no_border_cell(right_cell)

        left_cell.paragraphs[0].clear()
        for edu_entry in education:
            degree_parts = [edu_entry.get("studyType", ""), edu_entry.get("area", "")]
            degree = " - ".join(part for part in degree_parts if part)
            parts = [degree, edu_entry.get("institution", "")]
            text = " | ".join(part for part in parts if part)
            _cell_bullet(left_cell, text)

        right_cell.paragraphs[0].clear()
        for cert in certifications:
            _cell_bullet(right_cell, cert)
    else:
        for edu_entry in education:
            degree_parts = [edu_entry.get("studyType", ""), edu_entry.get("area", "")]
            degree = " - ".join(part for part in degree_parts if part)
            parts = [degree, edu_entry.get("institution", "")]
            text = " | ".join(part for part in parts if part)
            _bullet(doc, text, body_pt=int(theme.body_pt), font_name=theme.font)


def build_resume(resume_json: ResumeJSON, personal: dict, education: list,
                    output_path: str, theme: ThemeConfig = CLASSIC):
    """
    Build a resume document.
    """
    if theme.layout == "sidebar":
        return _build_resume_sidebar(resume_json, personal, education, output_path, theme)

    doc = Document()

    # Page margins from theme
    for section in doc.sections:
        section.top_margin = Inches(theme.margin_top)
        section.bottom_margin = Inches(theme.margin_bottom)
        section.left_margin = Inches(theme.margin_left)
        section.right_margin = Inches(theme.margin_right)

    # Compute right-tab position from actual page geometry (EMU → twips)
    current_section = doc.sections[0]
    page_width = current_section.page_width or 0
    left_margin = current_section.left_margin or 0
    right_margin = current_section.right_margin or 0
    text_width_emu = page_width - left_margin - right_margin
    right_tab_stop = int(text_width_emu / 635)  # 1 twip = 635 EMU

    _render_resume_header(doc, personal, theme)
    _section_heading(doc, "Professional Summary", theme)
    summary_text = resume_json.summary.strip()
    _body_paragraph(doc, summary_text, after_pt=2, body_pt=int(theme.body_pt),
                    font_name=theme.font)

    _render_resume_experience(doc, resume_json.experience, right_tab_stop, theme)
    _render_resume_skills(doc, resume_json.skill_categories, theme)

    if resume_json.projects:
        _render_resume_projects(doc, resume_json.projects,
                                resume_json.projects_section_heading, theme)

    _render_resume_education_certs(doc, education, resume_json.certifications, theme)

    _dirname = os.path.dirname(output_path)
    if _dirname:
        os.makedirs(_dirname, exist_ok=True)
    doc.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# Cover letter builder
# ---------------------------------------------------------------------------

COVER_PT = 12  # body font size for cover letters (larger than resume's 10pt)


def _cover_paragraph(doc: "DocxDocument", text: str, indent: bool = False,
                        before_pt: float = 6, after_pt: float = 6,
                        cover_pt: int = COVER_PT, font_name: str = "Arial") -> object:
    """Helper to add a formatted paragraph to the cover letter."""
    # pylint: disable=too-many-arguments,too-many-positional-arguments
    paragraph = doc.add_paragraph()
    _set_space(paragraph, before_pt=before_pt, after_pt=after_pt)
    if indent:
        paragraph.paragraph_format.first_line_indent = Inches(0.5)
    run = paragraph.add_run(text.strip())
    _set_font(run, cover_pt, font_name=font_name)
    return paragraph


def _cover_bullet(doc: "DocxDocument", text: str, cover_pt: int = COVER_PT,
                    font_name: str = "Arial"):
    bullet_paragraph = doc.add_paragraph(style="List Bullet")
    _set_space(bullet_paragraph, before_pt=0, after_pt=2)
    for run in bullet_paragraph.runs:
        bullet_paragraph._p.remove(run._r)  # pylint: disable=protected-access
    bullet_run = bullet_paragraph.add_run(text)
    _set_font(bullet_run, cover_pt, font_name=font_name)
    return bullet_paragraph


def _render_cover_letter_header(doc: "DocxDocument", personal: dict,
                                theme: ThemeConfig, cover_pt: int):
    # pylint: disable=too-many-locals
    name_paragraph = doc.add_paragraph()
    _set_space(name_paragraph, before_pt=0, after_pt=2)
    name_run = name_paragraph.add_run(personal["name"])
    _set_font(name_run, int(theme.name_pt), bold=True, color=RGBColor(*theme.accent_color),
                font_name=theme.font)

    loc = personal.get("location", {})
    loc_parts = [loc.get("city", ""), loc.get("region", ""), loc.get("countryCode", "")]
    location_str = (
        ", ".join(part for part in loc_parts if part)
        if isinstance(loc, dict) else str(loc or "")
    )
    linkedin = next(
        (p.get("username", "") for p in personal.get("profiles", [])
            if p.get("network") == "LinkedIn"), ""
    )
    github = next(
        (p.get("username", "") for p in personal.get("profiles", [])
            if p.get("network") == "GitHub"), ""
    )
    line1_parts = [
        location_str,
        personal.get("phone", ""),
        personal.get("email", ""),
    ]
    line2_parts = [
        f"LinkedIn: {linkedin}" if linkedin else "",
        f"GitHub: {github}" if github else "",
    ]
    contact_paragraph = doc.add_paragraph()
    _set_space(contact_paragraph, before_pt=0, after_pt=12)
    line1 = " \u2219 ".join(part for part in line1_parts if part)
    line2 = " \u2219 ".join(part for part in line2_parts if part)
    run1 = contact_paragraph.add_run(line1)
    _set_font(run1, cover_pt, font_name=theme.font)
    if line2:
        run1.add_break()
        run2 = contact_paragraph.add_run(line2)
        _set_font(run2, cover_pt, font_name=theme.font)


def _render_cover_letter_body(doc: "DocxDocument", cover_json: CoverLetterJSON, company: str,
                                theme: ThemeConfig, cover_pt: int):
    # Date (right-aligned)
    today = date.today().strftime("%B %d, %Y")
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_space(date_p, after_pt=12)
    date_run = date_p.add_run(today)
    _set_font(date_run, cover_pt, font_name=theme.font)

    # Greeting
    greeting_p = doc.add_paragraph()
    _set_space(greeting_p, after_pt=6)
    greeting_run = greeting_p.add_run(f"Dear Hiring Manager at {company},")
    _set_font(greeting_run, cover_pt, font_name=theme.font)

    _cover_paragraph(doc, cover_json.opening, indent=True, cover_pt=cover_pt, font_name=theme.font)
    for paragraph_text in cover_json.body_paragraphs:
        _cover_paragraph(doc, paragraph_text, indent=True, cover_pt=cover_pt, font_name=theme.font)

    if cover_json.highlights:
        intro_text = cover_json.highlights_intro or "A few highlights from my background:"
        _cover_paragraph(doc, intro_text, indent=True, before_pt=6, after_pt=2, cover_pt=cover_pt,
                            font_name=theme.font)
        for bullet_text in cover_json.highlights:
            _cover_bullet(doc, bullet_text, cover_pt=cover_pt, font_name=theme.font)

    _cover_paragraph(doc, cover_json.closing, indent=True, before_pt=8, cover_pt=cover_pt,
                        font_name=theme.font)


def build_cover_letter(cover_json: CoverLetterJSON, personal: dict,
                        company: str, output_path: str,
                        theme: ThemeConfig = CLASSIC):
    """
    Build a cover letter document.
    """
    # pylint: disable=too-many-arguments,too-many-positional-arguments,too-many-locals
    doc = Document()
    cover_pt = int(theme.body_pt) + 2

    for section in doc.sections:
        section.top_margin = Inches(max(theme.margin_top, 1.0))
        section.bottom_margin = Inches(max(theme.margin_bottom, 1.0))
        section.left_margin = Inches(max(theme.margin_left, 1.0))
        section.right_margin = Inches(max(theme.margin_right, 1.0))

    _render_cover_letter_header(doc, personal, theme, cover_pt)
    _render_cover_letter_body(doc, cover_json, company, theme, cover_pt)

    # Sign-off
    signoff_paragraph = doc.add_paragraph()
    _set_space(signoff_paragraph, before_pt=16, after_pt=2)
    signoff_run = signoff_paragraph.add_run("Sincerely,")
    _set_font(signoff_run, cover_pt, font_name=theme.font)

    sig_name_paragraph = doc.add_paragraph()
    _set_space(sig_name_paragraph, before_pt=24, after_pt=0)
    sig_name_run = sig_name_paragraph.add_run(personal["name"])
    _set_font(sig_name_run, cover_pt, bold=True, font_name=theme.font)

    sig_contact_parts = [personal.get("email", ""), personal.get("phone", "")]
    sig_contact_paragraph = doc.add_paragraph()
    _set_space(sig_contact_paragraph, before_pt=2, after_pt=0)
    sig_contact_line = "  |  ".join(part for part in sig_contact_parts if part)
    sig_contact_run = sig_contact_paragraph.add_run(sig_contact_line)
    _set_font(sig_contact_run, cover_pt, font_name=theme.font)

    _dirname = os.path.dirname(output_path)
    if _dirname:
        os.makedirs(_dirname, exist_ok=True)
    doc.save(output_path)
    return output_path
